from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


def flush_code_block(document: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    code_lines.clear()


def flush_table(document: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return
    rows: list[list[str]] = []
    for line in table_lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        table_lines.clear()
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            table.cell(row_index, col_index).text = text
    table_lines.clear()


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            paragraph.add_run(part)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    code_lines: list[str] = []
    table_lines: list[str] = []
    in_code_block = False

    for line in lines:
        if line.startswith("```"):
            flush_table(document, table_lines)
            if in_code_block:
                flush_code_block(document, code_lines)
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        flush_table(document, table_lines)

        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue

        if stripped == "---":
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.LINE)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1))
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered.group(2))
            continue

        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            paragraph = document.add_paragraph(style="Intense Quote")
            add_inline_runs(paragraph, quote.group(1))
            continue

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, line)

    flush_table(document, table_lines)
    flush_code_block(document, code_lines)
    document.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: md_to_docx.py <input.md> <output.docx>")
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
